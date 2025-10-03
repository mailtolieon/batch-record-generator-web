import pandas as pd
from docx import Document
import os
from datetime import datetime
import re
from pathlib import Path
from typing import List, Dict, Optional, Any
import logging

class WebDocumentGenerator:
    """
    Document generator adapted for web application use
    """
    
    def __init__(self, config):
        self.config = config
        self.df = None
        self.column_cache = {}
        self.setup_logging()
        self.load_excel_data()
    
    def setup_logging(self):
        """Setup logging for web application"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s'
        )
        self.logger = logging.getLogger(__name__)
    
    def load_excel_data(self) -> bool:
        """Load Excel data for web use"""
        try:
            if not os.path.exists(self.config.excel_file):
                self.logger.error(f"Excel file not found: {self.config.excel_file}")
                return False
            
            # Load Excel file
            excel_data = pd.ExcelFile(self.config.excel_file)
            
            # Validate sheet name
            if self.config.sheet_name not in excel_data.sheet_names:
                self.logger.warning(f"Sheet '{self.config.sheet_name}' not found. Using first sheet.")
                self.config.sheet_name = excel_data.sheet_names[0]
            
            # Load data
            self.df = pd.read_excel(
                self.config.excel_file,
                sheet_name=self.config.sheet_name,
                engine='openpyxl',
                dtype=str,
                na_values=['', 'NULL', 'null', 'NaN'],
                keep_default_na=False
            )
            
            # Clean column names
            self.df.columns = self.clean_column_names(self.df.columns)
            
            # Cache column mappings
            self.cache_column_mappings()
            
            self.logger.info(f"Loaded {len(self.df)} records")
            return True
            
        except Exception as e:
            self.logger.error(f"Error loading Excel: {e}")
            return False
    
    def clean_column_names(self, columns):
        """Clean column names for web use"""
        cleaned = []
        for col in columns:
            col_clean = re.sub(r'[^\w\s]', '', str(col))
            col_clean = col_clean.strip().lower().replace(' ', '_')
            col_clean = re.sub(r'_+', '_', col_clean)
            cleaned.append(col_clean)
        return cleaned
    
    def cache_column_mappings(self):
        """Cache column mappings for performance"""
        for field, possible_names in self.config.column_mappings.items():
            self.column_cache[field] = self.find_column_name(possible_names, field, silent=True)
    
    def find_column_name(self, possible_names, description="", silent=False):
        """Find column name with web-optimized logging"""
        if self.df is None or self.df.empty:
            return None
        
        for name in possible_names:
            if name in self.df.columns:
                if not silent:
                    self.logger.info(f"Found {description} column: '{name}'")
                return name
        
        # Try partial matching
        for col in self.df.columns:
            col_lower = col.lower()
            for name in possible_names:
                if name.lower() in col_lower:
                    if not silent:
                        self.logger.info(f"Found partial match for {description}: '{col}' contains '{name}'")
                    return col
        
        if not silent:
            self.logger.warning(f"Could not find {description} column")
        return None
    
    def get_unique_products(self) -> List[str]:
        """Get unique products for web display"""
        if self.df is None or self.df.empty:
            return []
        
        try:
            product_column = self.column_cache.get('product')
            if not product_column:
                return []
            
            unique_products = self.df[product_column].dropna()
            unique_products = unique_products[unique_products.astype(str).str.strip() != '']
            
            seen = set()
            result = []
            for product in unique_products:
                product_str = str(product).strip()
                if product_str and product_str not in seen:
                    seen.add(product_str)
                    result.append(product_str)
            
            result.sort()
            return result
            
        except Exception as e:
            self.logger.error(f"Error getting unique products: {e}")
            return []
    
    def search_product_batches(self, product_name: str) -> Optional[List[Dict]]:
        """Search for product batches - optimized for web"""
        if self.df is None or self.df.empty:
            return None
        
        try:
            product_column = self.column_cache.get('product')
            if not product_column:
                return None
            
            # Case-insensitive search
            mask = self.df[product_column].astype(str).str.lower() == product_name.lower()
            product_batches = self.df[mask]
            
            if not product_batches.empty:
                batches = self.process_batch_data(product_batches)
                sorted_batches = self.sort_batches(batches)
                return sorted_batches
            
            return None
            
        except Exception as e:
            self.logger.error(f"Error searching batches: {e}")
            return None
    
    def process_batch_data(self, product_batches):
        """Process batch data for web output"""
        batches = []
        
        for _, batch in product_batches.iterrows():
            batch_data = {
                'batch_no': self.get_formatted_value(batch, 'batch_no'),
                'mfg_date': self.format_date_properly(self.get_formatted_value(batch, 'mfg_date')),
                'expiry_date': self.format_date_properly(self.get_formatted_value(batch, 'expiry_date')),
                'total_batch_yield': self.format_percentage(self.get_formatted_value(batch, 'yield')),
                'total_batch_accountability': self.format_percentage(self.get_formatted_value(batch, 'accountability')),
                'location_rack_shelf': self.get_formatted_value(batch, 'location'),
                'remarks': self.get_formatted_value(batch, 'remarks'),
                'sent_to_document_room': self.format_date_properly(self.get_formatted_value(batch, 'sent_to_doc'))
            }
            batches.append(batch_data)
        
        return batches
    
    def get_formatted_value(self, batch, field):
        """Get formatted value for web display"""
        column_name = self.column_cache.get(field)
        if not column_name or column_name not in batch:
            return ''
        
        value = batch[column_name]
        if pd.isna(value) or value == '':
            return ''
        
        return str(value).strip()
    
    def sort_batches(self, batches):
        """Sort batches for web output"""
        try:
            def get_batch_key(batch):
                batch_no = batch.get('batch_no', '0')
                numbers = re.findall(r'\d+', batch_no)
                batch_num = int(numbers[0]) if numbers else 0
                return batch_num
            
            return sorted(batches, key=get_batch_key)
        except Exception as e:
            self.logger.error(f"Error sorting batches: {e}")
            return batches
    
    def format_date_properly(self, date_value):
        """Format dates for web display"""
        if not date_value or str(date_value).strip() == '':
            return ''
        
        date_str = str(date_value).strip().upper()
        if any(na_val in date_str for na_val in ['#N/A', '#NA', 'N/A', 'NA']):
            return date_str
        
        try:
            if hasattr(date_value, 'strftime'):
                return date_value.strftime('%d.%m.%Y')
            
            date_str = date_str.split(' ')[0]
            
            date_formats = [
                '%Y-%m-%d', '%d-%m-%Y', '%d/%m/%Y', '%m/%d/%Y',
                '%d.%m.%Y', '%Y.%m.%d'
            ]
            
            for fmt in date_formats:
                try:
                    parsed_date = datetime.strptime(date_str, fmt)
                    return parsed_date.strftime('%d.%m.%Y')
                except ValueError:
                    continue
            
            return date_str
            
        except Exception as e:
            self.logger.warning(f"Could not format date: {e}")
            return date_str
    
    def format_percentage(self, value):
        """Format percentages for web display"""
        if not value or str(value).strip() == '':
            return ''
        
        value_str = str(value).strip().upper()
        if any(na_val in value_str for na_val in ['#N/A', '#NA', 'N/A', 'NA']):
            return value_str
        
        try:
            clean_value = re.sub(r'[%\s]', '', value_str)
            if clean_value and clean_value != 'nan':
                num_value = float(clean_value)
                return f"{num_value:.2f}%"
            return ''
        except:
            return value_str
    
    def generate_single_document(self, product_name: str) -> Dict[str, Any]:
        """Generate single document - adapted for web use"""
        try:
            # Ensure output directory exists
            Path(self.config.output_folder).mkdir(exist_ok=True)
            
            # Search for product batches
            batches = self.search_product_batches(product_name)
            
            if not batches:
                return {
                    'success': False,
                    'error': f'No batches found for product: {product_name}',
                    'batch_count': 0
                }
            
            # Generate filename
            current_date = datetime.now().strftime("%Y-%m-%d")
            safe_name = re.sub(r'[^\w\s-]', '', product_name).strip()
            filename = f"{safe_name}_{current_date}.docx"
            filepath = os.path.join(self.config.output_folder, filename)
            
            # Check if template exists
            if not os.path.exists(self.config.template_file):
                return {
                    'success': False,
                    'error': f'Template file not found: {self.config.template_file}',
                    'batch_count': 0
                }
            
            # Load template and generate document
            doc = Document(self.config.template_file)
            
            # Replace product name placeholder
            self.fill_product_name_in_header(doc, product_name)
            
            # Fill batch table
            success = self.fill_batch_table_with_formatting(doc, batches)
            
            if success:
                doc.save(filepath)
                
                return {
                    'success': True,
                    'filename': filename,
                    'filepath': filepath,
                    'batch_count': len(batches),
                    'product': product_name
                }
            else:
                return {
                    'success': False,
                    'error': 'Failed to fill document template',
                    'batch_count': 0
                }
                
        except Exception as e:
            self.logger.error(f"Error generating document: {e}")
            return {
                'success': False,
                'error': str(e),
                'batch_count': 0
            }
    
    def fill_product_name_in_header(self, doc, product_name):
        """Fill product name in document header"""
        try:
            from docx.shared import Pt
            
            PLACEHOLDER = "{{PRODUCT_NAME}}"
            
            # Search in headers
            for section in doc.sections:
                headers = [
                    section.first_page_header,
                    section.even_page_header, 
                    section.header
                ]
                
                for header in headers:
                    if header:
                        # Check paragraphs
                        for paragraph in header.paragraphs:
                            if PLACEHOLDER in paragraph.text:
                                paragraph.text = paragraph.text.replace(PLACEHOLDER, product_name)
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.size = Pt(14)
                                return True
                        
                        # Check tables in headers
                        for table in header.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        if PLACEHOLDER in paragraph.text:
                                            paragraph.text = paragraph.text.replace(PLACEHOLDER, product_name)
                                            for run in paragraph.runs:
                                                run.bold = True
                                                run.font.size = Pt(14)
                                            return True
            
            # Search in main document
            for paragraph in doc.paragraphs:
                if PLACEHOLDER in paragraph.text:
                    paragraph.text = paragraph.text.replace(PLACEHOLDER, product_name)
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(14)
                    return True
            
            # Search in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if PLACEHOLDER in paragraph.text:
                                paragraph.text = paragraph.text.replace(PLACEHOLDER, product_name)
                                for run in paragraph.runs:
                                    run.bold = True
                                    run.font.size = Pt(14)
                                return True
            
            return False
            
        except Exception as e:
            self.logger.error(f"Error filling product name: {e}")
            return False
    
    def fill_batch_table_with_formatting(self, doc, batches):
        """Fill batch table with formatting"""
        try:
            if not doc.tables:
                return False
            
            table = doc.tables[0]
            
            # Clear existing data rows
            while len(table.rows) > 1:
                table._tbl.remove(table.rows[1]._tr)
            
            # Add new rows with data
            for batch in batches:
                row = table.add_row()
                cells = row.cells
                
                column_data = [
                    batch.get('batch_no', ''),
                    batch.get('mfg_date', ''),
                    batch.get('expiry_date', ''),
                    batch.get('total_batch_yield', ''),
                    batch.get('total_batch_accountability', ''),
                    batch.get('location_rack_shelf', ''),
                    batch.get('remarks', ''),
                    batch.get('sent_to_document_room', '')
                ]
                
                for col_idx, data in enumerate(column_data):
                    if col_idx < len(cells):
                        cells[col_idx].text = str(data) if data is not None else ''
                        for paragraph in cells[col_idx].paragraphs:
                            paragraph.alignment = 1  # Center alignment
            
            return True
            
        except Exception as e:
            self.logger.error(f"Error filling table: {e}")
            return False
    
    def preview_product_data(self, product_name: str) -> Optional[List[Dict]]:
        """Preview product data for web display"""
        return self.search_product_batches(product_name)